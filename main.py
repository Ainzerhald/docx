import datetime
import os
import re
import shutil

import win32com.client
from docx import Document
from prettytable import PrettyTable
from rich.progress import track

new_dir = "\\folder_for_tmp_docx_don't_touch\\"


def find_docx(path):
    files = []
    if not os.path.exists(path + new_dir):
        os.mkdir(path + new_dir)
    for file in track(os.listdir(path), description="Поиск файлов:"):
        if os.path.isfile(path + "\\" + file):
            mass = os.path.splitext(file)
            if mass[1] == ".docx":
                files.append(path + new_dir + file)
                shutil.copy(path + "\\" + file, path + new_dir + file)
            elif mass[1] == ".doc":
                word = win32com.client.Dispatch('Word.Application')
                wb = word.Documents.Open(path + "\\" + file)
                wb.SaveAs2(path + new_dir + mass[0] + ".docx", FileFormat=16)
                wb.Close()
                files.append(path + new_dir + mass[0] + ".docx")
    print("\n")
    return files


def get_data_from_files(files):
    fil = open(os.getcwd() + '\\errors.txt', 'w')
    tables = []
    for file in track(files, description="Чтение данных из файлов:"):
        try:
            document = Document(file)
            if len(document.tables) < 1:
                fil.write("Нет таблицы. Файл: " + os.path.basename(file) + "\n")
            else:
                tab = document.tables[0]
                key = []
                for index, row in enumerate(tab.rows):
                    deleted = False
                    text = []
                    for cell in row.cells:
                        if len(cell.text.strip()) > 0:
                            runs = cell.paragraphs[0].runs
                            color = "FFFFFF"
                            if len(runs) > 0:
                                color = runs[0].font.color.rgb
                            if str(color) == 'FF0000':
                                deleted = True
                                break
                            if index == 0:
                                key.append(cell.text.replace("\n", " "))
                            else:
                                text.append(cell.text.replace("\n", " "))
                    if not deleted:
                        dicts = dict(zip(key, text))
                        if len(dicts.keys()) > 0:
                            dicts["file"] = os.path.basename(file)
                            tables.append(dicts)
        except BaseException as e:
            print(e)
            fil.write("Ошибка чтения. Файл: " + os.path.basename(file) + "\n")
    fil.close()
    return tables


def isDate(date):
    try:
        date = re.sub("[^0-9.]", "", date.replace(' ', ''))
        if len(date) == 8:
            datetime.datetime.strptime(date, '%d.%m.%y').strftime('%d.%m.%Y')
        elif len(date) == 10:
            datetime.datetime.strptime(date, '%d.%m.%Y').strftime('%d.%m.%Y')
        else:
            return False
        return True
    except BaseException:
        return False


start_date_string = ""
start_date = datetime.date
day = month = year = ""
error_date = True
while error_date:
    start_date_string = input("Введите месяц и год: ")
    if start_date_string.replace(" ", "") == "":
        start_date = start_date.today()
        break
    if len(start_date_string) != 7:
        print("Некорректная дата. Пример: 01.2021")
    else:
        month, year = start_date_string.split(".")
        try:
            start_date = datetime.date(int(year), int(month), 1)
            error_date = False
        except BaseException:
            day = month = year = ""
            print("Некорректная дата. Пример: 01.2021")

table = PrettyTable()
os.chdir('C:\\Users\\Andrey\\Desktop\\test')
rows = get_data_from_files(find_docx(os.getcwd()))
data_for_table = []
unsorted_list = []
table.field_names = ["Гос.Номер", "Оплата", "Файл"]
for row in rows:
    nomer = ""
    oplata = ""
    for i in row:
        if i.lower().find("гос") != -1:
            nomer = row[i]
        if i.lower().find("оплата") != -1:
            oplata = row[i]
    if bool(nomer) and bool(oplata):
        if isDate(oplata):
            oplata = re.sub("[^0-9.]", "", oplata.replace(' ', ''))
            if len(oplata) == 8:
                oplata = datetime.datetime.strptime(oplata, '%d.%m.%y').strftime('%d.%m.%Y')
            elif len(oplata) == 10:
                oplata = datetime.datetime.strptime(oplata, '%d.%m.%Y').strftime('%d.%m.%Y')
            unsorted_list.append([nomer, oplata, row["file"]])
        else:
            fil = open(os.getcwd() + '\\errors.txt', 'a')
            fil.write("Ошибка в дате. Дата: " + oplata + " Файл: " + row["file"] + "\n")
            fil.close()
'''sorted_list = sorted(unsorted_list,
                     key=lambda date: datetime.datetime.toordinal(datetime.datetime.strptime(date[1], '%d.%m.%Y')))'''
month_str = ""
year_str = ""
file_name = ""
change = False
for i in unsorted_list:
    day, month, year = i[1].split(".")
    checked_date = datetime.date(int(year), int(month), int(day))
    '''if month_str != checked_date.strftime("%B"):
        month_str = checked_date.strftime("%B")
        change = True
    if year_str != checked_date.strftime("%Y"):
        year_str = checked_date.strftime("%Y")
        change = True'''
    if file_name != i[2]:
        file_name = i[2]
        change = True
    if start_date.year == checked_date.year and start_date.month == checked_date.month:
        if change:
            table.add_row(["", "", ""])
            '''table.add_row(["", month_str, year_str])
            table.add_row(["", "", ""])'''
            change = False
        table.add_row(i)
fil = open(os.getcwd() + '\\result.txt', 'w')
fil.write(str(table))
fil.close()
shutil.rmtree(os.getcwd() + new_dir)
